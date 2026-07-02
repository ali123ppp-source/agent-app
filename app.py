import streamlit as st
import pandas as pd
from io import BytesIO
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import RGBColor, Pt, Inches
import re
from datetime import datetime

# =============================================================================
# إعدادات واجهة المستخدم وتنسيقات الـ CSS للويب
# =============================================================================
st.set_page_config(page_title="نظام المقارنة المتطور للوكلاء", layout="wide")
st.markdown("""
    <style>
    th, td { text-align: right !important; dir: rtl !important; white-space: nowrap !important; }
    div.stButton > button { background-color: #2C3E50; color: white; width: 100%; font-weight: bold; border-radius: 8px;}
    .report-box { background-color: #ECF0F1; padding: 15px; border-radius: 8px; border-right: 5px solid #2C3E50; text-align: right; margin-bottom: 10px;}
    div[data-testid="stRadio"] > label { font-weight: bold; color: #2C3E50; font-size: 16px; }
    .date-badge { display: inline-block; padding: 8px 12px; background-color: #F8F9F9; color: #2C3E50; border-radius: 5px; font-weight: bold; font-size: 15px; border: 1px solid #BDC3C7; margin-bottom: 5px; direction: rtl; width: 100%; text-align: center;}
    .date-badge span.old { color: #C0392B; }
    .date-badge span.new { color: #27AE60; }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 style='text-align: right;'>نظام المقارنة الشامل والذكي 📄🔎</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: right;'>تمت إضافة ميزة الترقيم التلقائي لصفحات الوورد والتعرف الذكي على التواريخ في الترويسة السفلية.</p>", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 1. محرك الاستشعار الزمني (محسن لاكتشاف التاريخ في الترويسة السفلية)
# -----------------------------------------------------------------------------
def extract_document_date(doc):
    patterns = [
        r"([A-Za-z]+,\s+[A-Za-z]+\s+\d{1,2},\s+\d{4})", # Sunday, May 17, 2026
        r"(\d{1,2}[/-]\d{1,2}[/-]\d{4})",               # 17/05/2026 أو 17-05-2026
        r"(\d{4}[/-]\d{1,2}[/-]\d{1,2})",               # 2026/05/17
        r"(\d{1,2}\s+[\u0600-\u06FF]+\s+\d{4})"         # تواريخ عربية مثل: 15 أيار 2024
    ]
    
    # البحث في الترويسة السفلية (Footer) أولاً
    for section in doc.sections:
        if section.footer:
            for para in reversed(section.footer.paragraphs):
                for pattern in patterns:
                    match = re.search(pattern, para.text)
                    if match:
                        try:
                            d_str = match.group(1)
                            if "-" in d_str or "/" in d_str: return pd.to_datetime(d_str, dayfirst=True).to_pydatetime()
                            if re.search(r"[\u0600-\u06FF]", d_str): return datetime.now() # حالة تقديرية للعربي
                            return datetime.strptime(d_str, "%A, %B %d, %Y")
                        except: continue
                        
    # البحث في آخر 50 سطراً
    paragraphs = doc.paragraphs[-50:] if len(doc.paragraphs) > 50 else doc.paragraphs
    for para in reversed(paragraphs):
        for pattern in patterns:
            match = re.search(pattern, para.text)
            if match:
                try:
                    d_str = match.group(1)
                    if "-" in d_str or "/" in d_str: return pd.to_datetime(d_str, dayfirst=True).to_pydatetime()
                    if re.search(r"[\u0600-\u06FF]", d_str): return datetime.now()
                    return datetime.strptime(d_str, "%A, %B %d, %Y")
                except: continue
    
    # الطبقة الثالثة: خصائص الملف (Metadata)
    try:
        if doc.core_properties.modified:
            return doc.core_properties.modified.replace(tzinfo=None)
    except:
        pass
        
    return None

# -----------------------------------------------------------------------------
# 2. محرك الاستخراج الدقيق
# -----------------------------------------------------------------------------
def extract_clean_records(doc, card_type="old"):
    records = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        cells = [c.strip() for c in text.split(',')]
        if len(cells) >= 6 and any(char.isdigit() for char in cells[0]) and any('\u0600' <= char <= '\u06FF' for char in cells[3]):
            try:
                withheld, eligible, total, name = int(cells[0]), int(cells[1]), int(cells[2]), cells[3]
                old_card = cells[4]
                new_card = cells[5] if len(cells) > 5 else old_card
                selected_card = old_card if card_type == "old" else new_card
                seq = cells[6] if len(cells) > 6 else "-"
                if selected_card:
                    records[selected_card] = {"seq": seq, "name": name, "total": total, "eligible": eligible, "withheld": withheld}
            except ValueError: continue

    if not records:
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                if not any(cells) or "المركز" in "".join(cells) or "الوكيل" in "".join(cells) or "اسم رب" in "".join(cells): continue
                name_idx = -1
                max_len = 0
                for i, c in enumerate(cells):
                    if any('\u0600' <= char <= '\u06FF' for char in c) and not any(char.isdigit() for char in c):
                        if len(c) > max_len: max_len, name_idx = len(c), i
                if name_idx == -1: continue
                card_indices = [i for i, c in enumerate(cells) if c.isdigit() and len(c) >= 5]
                if not card_indices: continue
                
                old_card = cells[card_indices[0]]
                new_card = cells[card_indices[-1]] if len(card_indices) > 1 else old_card
                selected_card = old_card if card_type == "old" else new_card
                seq = "-"
                for i in range(len(cells)-1, card_indices[-1], -1):
                    if cells[i].isdigit():
                        seq = cells[i]
                        break
                digit_cells = [int(cells[i]) for i in range(name_idx) if cells[i].isdigit()]
                if len(digit_cells) >= 3: withheld, eligible, total = digit_cells[0], digit_cells[1], digit_cells[2]
                elif len(digit_cells) == 2: withheld, eligible, total = 0, digit_cells[0], digit_cells[1]
                else: continue
                records[selected_card] = {"seq": seq, "name": cells[name_idx], "total": total, "eligible": eligible, "withheld": withheld}
    return records

# -----------------------------------------------------------------------------
# 3. محرك المقارنة ونظام الإحالة الذكي
# -----------------------------------------------------------------------------
def process_comparison(old_data, new_data, mode, card_col_name, matching_engine):
    results = []
    results_type_1_reference = []
    counters = {
        "total_fam": 0, "eligible_fam": 0, "withheld_fam": 0, "added_fam": 0, "deleted_fam": 0,
        "inc_total": 0, "dec_total": 0, "net_total": 0, "inc_eligible": 0, "dec_eligible": 0, "net_eligible": 0,
        "inc_withheld": 0, "dec_withheld": 0, "net_withheld": 0
    }
    
    skip_seq_matching = (matching_engine == "محرك تخطي التسلسل (بطاقة فقط)")
    all_cards = set(old_data.keys()).union(set(new_data.keys()))
    
    for card in all_cards:
        if card in old_data and card in new_data:
            old_v, new_v = old_data[card], new_data[card]
            d_tot = new_v["total"] - old_v["total"]
            d_elig = new_v["eligible"] - old_v["eligible"]
            d_with = new_v["withheld"] - old_v["withheld"]
            is_changed = d_tot != 0 or d_elig != 0 or d_with != 0
            
            target_seq = new_v["seq"] if skip_seq_matching else old_v["seq"]
            notes = []
            
            if old_v["name"] != new_v["name"]:
                notes.append(f"تم تغيير الاسم / السابق / {old_v['name']}")
                is_changed = True
            
            if new_v["withheld"] == new_v["total"] and new_v["total"] > 0 and d_with > 0:
                notes.append("حجب كلي ❌")
            else:
                if d_with > 0: notes.append(f"تم حجب {d_with} نفر ➖")
                elif d_with < 0: notes.append(f"تم رفع الحجب عن {abs(d_with)} نفر ➕")
            
            if d_tot > 0: notes.append("إضافة طفل 👶")
            elif d_tot < 0: notes.append(f"نقصان {abs(d_tot)} نفر")
            
            referral_text = " | ".join(notes) if notes else ("تحديث بيانات" if is_changed else "")
            
            if is_changed:
                if d_tot != 0:
                    counters["total_fam"] += 1
                    counters["net_total"] += d_tot
                    if d_tot > 0: counters["inc_total"] += d_tot
                    else: counters["dec_total"] += abs(d_tot)
                if d_elig != 0:
                    counters["eligible_fam"] += 1
                    counters["net_eligible"] += d_elig
                    if d_elig > 0: counters["inc_eligible"] += d_elig
                    else: counters["dec_eligible"] += abs(d_elig)
                if d_with != 0:
                    counters["withheld_fam"] += 1
                    counters["net_withheld"] += d_with
                    if d_with > 0: counters["inc_withheld"] += d_with
                    else: counters["dec_withheld"] += abs(d_with)
                
                base_dict = {
                    "التسلسل": target_seq, "اسم رب الأسرة": new_v["name"], card_col_name: card,
                    "الأفراد الكلية": new_v["total"], "الأفراد المستحقة": new_v["eligible"], 
                    "الأفراد المحجوبين": new_v["withheld"], "الإحالة": referral_text, "meta_card": card
                }
                
                results_type_1_reference.append({**base_dict, "meta_status": "modified"})
                
                if mode == "النوع الأول" or mode == "النوع الثالث":
                    results.append({**base_dict, "meta_status": "modified", "meta_sort": 1})
                elif mode == "النوع الثاني":
                    results.append({
                        "التسلسل": target_seq, "اسم رب الأسرة": old_v["name"], card_col_name: card, "الحالة": "السابق",
                        "الأفراد الكلية": old_v["total"], "الأفراد المستحقة": old_v["eligible"], "الأفراد المحجوبين": old_v["withheld"],
                        "الإحالة": "", "meta_status": "type2_old", "meta_card": card, "meta_sort": 1
                    })
                    results.append({
                        "التسلسل": target_seq, "اسم رب الأسرة": new_v["name"], card_col_name: card, "الحالة": "الحديث",
                        "الأفراد الكلية": new_v["total"], "الأفراد المستحقة": new_v["eligible"], "الأفراد المحجوبين": new_v["withheld"],
                        "الإحالة": referral_text, "meta_status": "type2_new", "meta_card": card, "meta_sort": 2
                    })
            elif mode == "النوع الثالث":
                results.append({
                    "التسلسل": target_seq, "اسم رب الأسرة": old_v["name"], card_col_name: card,
                    "الأفراد الكلية": new_v["total"], "الأفراد المستحقة": new_v["eligible"], 
                    "الأفراد المحجوبين": new_v["withheld"], "الإحالة": "", "meta_status": "normal", "meta_card": card, "meta_sort": 1
                })
                
        elif card in old_data and card not in new_data:
            old_v = old_data[card]
            counters["deleted_fam"] += 1
            counters["dec_total"] += old_v["total"]
            counters["net_total"] -= old_v["total"]
            counters["dec_eligible"] += old_v["eligible"]
            counters["net_eligible"] -= old_v["eligible"]
            counters["dec_withheld"] += old_v["withheld"]
            counters["net_withheld"] -= old_v["withheld"]
            
            base_row = {"التسلسل": old_v["seq"], "اسم رب الأسرة": old_v["name"], card_col_name: card,
                        "الأفراد الكلية": old_v["total"], "الأفراد المستحقة": old_v["eligible"], 
                        "الأفراد المحجوبين": old_v["withheld"], "الإحالة": "عائلة محذوفة ❌", "meta_card": card}
            results_type_1_reference.append({**base_row, "meta_status": "deleted"})
            
            if mode == "النوع الثاني":
                results.append({**base_row, "الحالة": "محذوف", "meta_status": "deleted", "meta_card": card, "meta_sort": 1})
            else:
                results.append({**base_row, "meta_status": "deleted", "meta_sort": 1})
                
        elif card not in old_data and card in new_data:
            new_v = new_data[card]
            counters["added_fam"] += 1
            counters["inc_total"] += new_v["total"]
            counters["net_total"] += new_v["total"]
            counters["inc_eligible"] += new_v["eligible"]
            counters["net_eligible"] += new_v["eligible"]
            counters["inc_withheld"] += new_v["withheld"]
            counters["net_withheld"] += new_v["withheld"]
            
            base_row = {"التسلسل": new_v["seq"], "اسم رب الأسرة": new_v["name"], card_col_name: card,
                        "الأفراد الكلية": new_v["total"], "الأفراد المستحقة": new_v["eligible"], 
                        "الأفراد المحجوبين": new_v["withheld"], "الإحالة": "عائلة مضافة ✨", "meta_card": card}
            results_type_1_reference.append({**base_row, "meta_status": "added"})
            
            if mode == "النوع الثاني":
                results.append({**base_row, "الحالة": "مضاف", "meta_status": "added", "meta_card": card, "meta_sort": 1})
            else:
                results.append({**base_row, "meta_status": "added", "meta_sort": 1})
                
    return results, results_type_1_reference, counters

# -----------------------------------------------------------------------------
# 4. دوال التظليل البصري للويب
# -----------------------------------------------------------------------------
def style_all_types(doc_df, old_data, new_data, card_col_name, mode):
    styles = pd.DataFrame('', index=doc_df.index, columns=doc_df.columns)
    for idx, row in doc_df.iterrows():
        status = row.get("meta_status", "")
        card = row.get("meta_card")
        notes = str(row.get("الإحالة", ""))
        
        if "تم تغيير الاسم" in notes: styles.loc[idx, "الإحالة"] = 'color: #2980B9; font-weight: bold;'
        elif "إضافة طفل" in notes: styles.loc[idx, "الإحالة"] = 'color: #1ABC9C; font-weight: bold;'
        elif "حجب كلي" in notes or "تم حجب" in notes: styles.loc[idx, "الإحالة"] = 'color: #C0392B; font-weight: bold;'
        elif "تم رفع الحجب" in notes or "مضافة" in notes: styles.loc[idx, "الإحالة"] = 'color: #27AE60; font-weight: bold;'
        
        if status == "type2_old": styles.loc[idx, "الحالة"] = 'background-color: #F5F5F5; font-weight: bold; color: #7F8C8D;'
        elif status == "type2_new": styles.loc[idx, "الحالة"] = 'background-color: #E8F8F5; font-weight: bold; color: #16A085;'
        elif status == "added": styles.loc[idx] = 'background-color: #E8F5E9; color: #2E7D32;'
        elif status == "deleted": styles.loc[idx] = 'background-color: #ECEFF1; color: #455A64; text-decoration: line-through;'

        if status in ["modified", "type2_old", "type2_new"] and card in old_data and card in new_data:
            o_val, n_val = old_data[card], new_data[card]
            if o_val["total"] != n_val["total"]: styles.loc[idx, "الأفراد الكلية"] = 'background-color: #FDE0DC; font-weight: bold; color: #C0392B;'
            if o_val["eligible"] != n_val["eligible"]: styles.loc[idx, "الأفراد المستحقة"] = 'background-color: #FDE0DC; font-weight: bold; color: #C0392B;'
            if o_val["withheld"] != n_val["withheld"]: styles.loc[idx, "الأفراد المحجوبين"] = 'background-color: #FDE0DC; font-weight: bold; color: #C0392B;'
    return styles

# -----------------------------------------------------------------------------
# 5. محرك التصدير الفخم مع الترقيم التلقائي
# -----------------------------------------------------------------------------
def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} fill="{color_hex}"/>')
    tcPr.append(shd)

def clean_to_triple_name(name_str):
    if not name_str or pd.isna(name_str): return ""
    words = str(name_str).strip().split()
    return " ".join(words[:3])

def format_run(run, font_name="Microsoft Sans Serif", size_pt=14, color_rgb=None, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_rgb: run.font.color.rgb = color_rgb
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)

# دالة الترقيم التلقائي للصفحات
def add_page_number(paragraph):
    p = paragraph._p
    
    run_text = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "الصفحة "
    run_text.append(t)
    p.append(run_text)

    run_fld = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run_fld.append(fldChar1)
    run_fld.append(instrText)
    run_fld.append(fldChar2)
    run_fld.append(fldChar3)
    p.append(run_fld)

def create_word_table_report(doc_df, title, mode, card_col_name, old_data, new_data, new_file_name):
    doc = Document()
    
    for section in doc.sections:
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        w, h = section.page_height, section.page_width
        section.page_width = w
        section.page_height = h
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # تفعيل الترقيم في الترويسة السفلية للورقة
        footer = section.footer
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(p_footer)
        
    def append_table_to_doc(target_doc, df_to_write, table_title):
        agent_name = new_file_name.replace(".docx", "")
        agent_name = re.sub(r'(FOOD|FLOUR)', '', agent_name, flags=re.IGNORECASE)
        agent_name = agent_name.strip("- ").strip()
        
        banner_table = target_doc.add_table(rows=1, cols=1)
        banner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        banner_cell = banner_table.rows[0].cells[0]
        set_cell_background(banner_cell, "1ABC9C")
        tcPr = banner_cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="180" w:type="dxa"/><w:bottom w:w="180" w:type="dxa"/><w:left w:w="250" w:type="dxa"/><w:right w:w="250" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)
        p_banner = banner_cell.paragraphs[0]
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_b1 = p_banner.add_run(f"تقرير متغيرات الوكيل: {agent_name}")
        format_run(run_b1, font_name="Microsoft Sans Serif", size_pt=16, color_rgb=RGBColor(255, 255, 255), bold=True)
        
        if "FOOD" in new_file_name.upper():
            run_b2 = p_banner.add_run(" (غذائية)")
            format_run(run_b2, font_name="Microsoft Sans Serif", size_pt=16, color_rgb=RGBColor(255, 243, 79), bold=True)
        elif "FLOUR" in new_file_name.upper():
            run_b2 = p_banner.add_run(" (طحين)")
            format_run(run_b2, font_name="Microsoft Sans Serif", size_pt=16, color_rgb=RGBColor(240, 254, 240), bold=True)
            
        p_sub = target_doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(f"\n{table_title}")
        format_run(run_sub, font_name="Microsoft Sans Serif", size_pt=13, color_rgb=RGBColor(44, 62, 80), bold=True)
        
        headers = ["ت", "اسم المواطن", "التسلسل القديم", "الكلي", "المستحق", "المحجوب", "الحالة"]
        col_widths = [Inches(0.6), Inches(3.4), Inches(1.2), Inches(0.8), Inches(0.8), Inches(0.8), Inches(2.2)]
        
        table = target_doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = table._element.tblPr
        bidiVisual = parse_xml(f'<w:bidiVisual {nsdecls("w")}/>')
        tblPr.append(bidiVisual)
        
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            hdr_cells[i].width = col_widths[i]
            set_cell_background(hdr_cells[i], "E8ECEF")
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs: format_run(p.runs[0], font_name="Microsoft Sans Serif", size_pt=12, color_rgb=RGBColor(44, 62, 80), bold=True)
        
        for row_idx, (_, row) in enumerate(df_to_write.iterrows()):
            row_cells = table.add_row().cells
            row_bg = "FFFFFF" if row_idx % 2 == 0 else "F8F9F9"
            
            val_seq_new = str(row_idx + 1)
            val_name = clean_to_triple_name(row.get("اسم رب الأسرة", ""))
            val_seq_old = str(row.get("التسلسل", ""))
            val_total = str(row.get("الأفراد الكلية", ""))
            val_eligible = str(row.get("الأفراد المستحقة", ""))
            val_withheld = str(row.get("الأفراد المحجوبين", ""))
            
            ref_text = ""
            if "الإحالة" in row and pd.notna(row["الإحالة"]) and str(row["الإحالة"]).strip() != "": ref_text = str(row["الإحالة"]).strip()
            elif "الحالة" in row and pd.notna(row["الحالة"]) and str(row["الحالة"]).strip() != "": ref_text = str(row["الحالة"]).strip()
                
            cell_values = [val_seq_new, val_name, val_seq_old, val_total, val_eligible, val_withheld, ref_text]
            
            for i in range(len(headers)):
                cell = row_cells[i]
                cell.width = col_widths[i]
                cell.text = cell_values[i]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_background(cell, row_bg)
                
                if i == 0: set_cell_background(cell, "E5E7E9")
                elif i == 2: set_cell_background(cell, "FADBD8")
                    
                if p.runs:
                    run = p.runs[0]
                    color_rgb, bold_flag = None, False
                    
                    if i == 3: color_rgb, bold_flag = RGBColor(0, 51, 204), True
                    elif i == 4: color_rgb, bold_flag = RGBColor(0, 128, 0), True
                    elif i == 5: color_rgb, bold_flag = RGBColor(204, 0, 0), True
                    elif i == 6: color_rgb, bold_flag = RGBColor(102, 0, 153), True
                    else: color_rgb = RGBColor(0, 0, 0)
                        
                    format_run(run, font_name="Microsoft Sans Serif", size_pt=14, color_rgb=color_rgb, bold=bold_flag)
                    
        target_doc.add_paragraph()

    append_table_to_doc(doc, doc_df, title)
    
    cases_to_extract = [
        ("حالات تغيير اسم رب الأسرة", "تم تغيير الاسم"), ("حالات إضافة طفل", "إضافة طفل"),
        ("حالات حجب كلي", "حجب كلي"), ("حالات حجب نفر", "تم حجب"), ("حالات رفع الحجب", "تم رفع الحجب"),
        ("العوائل المضافة", "عائلة مضافة"), ("العوائل المحذوفة", "عائلة محذوفة")
    ]
    
    for case_title, keyword in cases_to_extract:
        if keyword == "تم حجب": matched_mask = doc_df['الإحالة'].str.contains("تم حجب", na=False) & ~doc_df['الإحالة'].str.contains("حجب كلي", na=False)
        else: matched_mask = doc_df['الإحالة'].str.contains(keyword, na=False)
            
        if mode == "النوع الثاني":
            matched_cards = doc_df[matched_mask]['meta_card'].dropna().unique()
            case_df = doc_df[doc_df['meta_card'].isin(matched_cards)]
        else: case_df = doc_df[matched_mask]
            
        if not case_df.empty:
            doc.add_page_break()
            append_table_to_doc(doc, case_df, f"كشف مستقل: {case_title}")
                
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_word_stats_report(counters, filename_base):
    doc = Document()
    doc.add_heading(f"تقرير الإحصاء - للملف: {filename_base}", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run("أولاً: إحصاء حركة الأفراد").bold = True
    stats_individuals = [
        ("زيادة الكلية:", f"+{counters['inc_total']}"), ("نقصان الكلية:", f"-{counters['dec_total']}"), ("صافي الكلية:", f"{counters['net_total']:+d}"),
        ("زيادة المستحقة:", f"+{counters['inc_eligible']}"), ("نقصان المستحقة:", f"-{counters['dec_eligible']}"), ("صافي المستحقة:", f"{counters['net_eligible']:+d}"),
        ("زيادة المحجوبين:", f"+{counters['inc_withheld']}"), ("نقصان المحجوبين:", f"-{counters['dec_withheld']}"), ("صافي المحجوبين:", f"{counters['net_withheld']:+d}")
    ]
    for text, val in stats_individuals: doc.add_paragraph().add_run(f"{val} : {text}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().add_run("ثانياً: العوائل").bold = True
    stats_families = [
        ("تغيرت الكلية:", counters['total_fam']), ("تغيرت المستحقة:", counters['eligible_fam']),
        ("تغيرت المحجوبين:", counters['withheld_fam']), ("عوائل مضافة:", counters['added_fam']), ("عوائل محذوفة:", counters['deleted_fam'])
    ]
    for text, val in stats_families: doc.add_paragraph().add_run(f"{val} : {text}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# 6. الواجهة الرئيسية
# -----------------------------------------------------------------------------
st.markdown("<h3 style='text-align: right;'>📂 منطقة الرفع والمطابقة</h3>", unsafe_allow_html=True)
uploaded_files = st.file_uploader("ارفع ملفي الشهر السابق والحالي معاً", type=['docx'], accept_multiple_files=True)

col_opts1, col_opts2, col_opts3 = st.columns(3)
with col_opts1: comparison_mode = st.radio("🎯 نوع المقارنة:", ["النوع الأول", "النوع الثاني", "النوع الثالث"], horizontal=True)
with col_opts2: card_choice_ui = st.radio("💳 البطاقة المعتمدة:", ["رقم البطاقة القديم", "رقم البطاقة الحديث"], horizontal=True)
with col_opts3: matching_engine = st.radio("⚙️ محرك المطابقة المستهدف:", ["المحرك القياسي", "محرك تخطي التسلسل (بطاقة فقط)"], horizontal=True)

card_type_param = "old" if card_choice_ui == "رقم البطاقة القديم" else "new"
card_col_name = card_choice_ui
swap_files = st.checkbox("🔄 **عكس الملفين يدوياً (القديم يصبح حديثاً والحديث قديماً)**")

grid_column_configuration = {
    "التسلسل": st.column_config.TextColumn("التسلسل", width="small"),
    "اسم رب الأسرة": st.column_config.TextColumn("اسم رب الأسرة", width="large"),
    card_col_name: st.column_config.TextColumn(card_col_name, width="medium"),
    "الحالة": st.column_config.TextColumn("الحالة", width="small"),
    "الأفراد الكلية": st.column_config.NumberColumn("الأفراد الكلية", width="small"),
    "الأفراد المستحقة": st.column_config.NumberColumn("الأفراد المستحقة", width="small"),
    "الأفراد المحجوبين": st.column_config.NumberColumn("الأفراد المحجوبين", width="small"),
    "الإحالة": st.column_config.TextColumn("الإحالة", width="large")
}

if st.button("بدء المقارنة الذكية واستخراج المتغيرات والأوراق"):
    if len(uploaded_files) == 2:
        with st.spinner('جاري التحليل وعزل الحالات تلقائياً...'):
            doc1, doc2 = Document(uploaded_files[0]), Document(uploaded_files[1])
            date1, date2 = extract_document_date(doc1), extract_document_date(doc2)
            
            file_a_is_older = (date1 < date2) if (date1 and date2) else True
            if swap_files: file_a_is_older = not file_a_is_older
                
            if file_a_is_older:
                old_doc, new_doc = doc1, doc2
                old_name, new_name = uploaded_files[0].name, uploaded_files[1].name
            else:
                old_doc, new_doc = doc2, doc1
                old_name, new_name = uploaded_files[1].name, uploaded_files[0].name

            st.markdown(f"<div class='date-badge'>الملف المعتمد كـ <span class='old'>السابق: ({old_name})</span> | الملف المعتمد كـ <span class='new'>الحديث: ({new_name})</span></div>", unsafe_allow_html=True)
            
            old_data = extract_clean_records(old_doc, card_type=card_type_param)
            new_data = extract_clean_records(new_doc, card_type=card_type_param)
            
            results, results_ref, counters = process_comparison(old_data, new_data, comparison_mode, card_col_name, matching_engine)
            
            if results:
                results = sorted(results, key=lambda x: (str(x.get("اسم رب الأسرة", "")), x.get("meta_sort", 0)))
                results_ref = sorted(results_ref, key=lambda x: str(x.get("اسم رب الأسرة", "")))
                
                df_results = pd.DataFrame(results)
                df_results_full = df_results.copy()
                df_display = df_results.copy()
                
                if comparison_mode == "النوع الثاني":
                    for idx, row in df_display.iterrows():
                        if row.get("meta_status") == "type2_new":
                            df_display.at[idx, "التسلسل"], df_display.at[idx, "اسم رب الأسرة"], df_display.at[idx, card_col_name], df_display.at[idx, "الإحالة"] = "", "", "", ""
                
                st.markdown(f"<h3 style='text-align: right;'>📋 المخرجات الشاشاتية ({comparison_mode})</h3>", unsafe_allow_html=True)
                
                styled_df = df_display.style.apply(lambda d: style_all_types(d, old_data, new_data, card_col_name, comparison_mode), axis=None)
                if comparison_mode == "النوع الثاني": cols_order = ["التسلسل", "اسم رب الأسرة", card_col_name, "الحالة", "الأفراد الكلية", "الأفراد المستحقة", "الأفراد المحجوبين", "الإحالة"]
                else: cols_order = ["التسلسل", "اسم رب الأسرة", card_col_name, "الأفراد الكلية", "الأفراد المستحقة", "الأفراد المحجوبين", "الإحالة"]
                
                st.dataframe(styled_df, use_container_width=True, hide_index=True, column_order=cols_order, column_config=grid_column_configuration)
                
                base_name = new_name.rsplit('.', 1)[0]
                col_dl1, col_dl2 = st.columns(2)
                with col_dl1:
                    word_report = create_word_table_report(df_results_full, f"تقرير - {comparison_mode}", comparison_mode, card_col_name, old_data, new_data, new_name)
                    st.download_button(label="📥 تحميل المخرجات Word بالديكور الجديد الفخم", data=word_report, file_name=f"تقرير_{base_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                with col_dl2:
                    word_stats = create_word_stats_report(counters, base_name)
                    st.download_button(label="📊 تحميل تقرير الإحصاء Word", data=word_stats, file_name=f"احصائيات_{base_name}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.success("🎉 تطابق تام! لا توجد فروقات بين الملفين.")
    else:
        st.warning("⚠️ يرجى رفع ملفين اثنين بالضبط للتمكن من بدء المقارنة.")
